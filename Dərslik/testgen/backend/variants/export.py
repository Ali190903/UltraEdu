import io
import json

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Flowable
from docx import Document
from docx.shared import Pt, Mm

from .latex_render import latex_to_image, split_text_and_latex


def export_json(variant_data: dict) -> bytes:
    questions = []
    for item in variant_data["questions"]:
        q = item["question"]
        questions.append({
            "order": item["order"],
            "question_text": q.question_text,
            "options": q.options,
            "correct_answer": q.correct_answer,
            "explanation": q.explanation,
            "difficulty": q.difficulty,
            "topic": q.topic,
        })

    result = {
        "title": variant_data["variant"].title,
        "subject": variant_data["variant"].subject,
        "total_questions": variant_data["variant"].total_questions,
        "questions": questions,
    }
    return json.dumps(result, ensure_ascii=False, indent=2).encode("utf-8")


def export_text(variant_data: dict) -> str:
    lines = [
        f"# {variant_data['variant'].title}",
        f"Fənn: {variant_data['variant'].subject}",
        f"Sual sayı: {variant_data['variant'].total_questions}",
        "",
    ]
    for item in variant_data["questions"]:
        q = item["question"]
        lines.append(f"**{item['order']}.** {q.question_text}")
        if q.options:
            for key, val in q.options.items():
                lines.append(f"   {key}) {val}")
        lines.append("")

    lines.append("\n--- Cavablar ---\n")
    for item in variant_data["questions"]:
        q = item["question"]
        lines.append(f"{item['order']}. {q.correct_answer}")

    return "\n".join(lines)


class InlineLatexImage(Flowable):
    """A flowable that renders inline with text height."""

    def __init__(self, img_bytes: bytes, max_height: float = 5 * mm):
        super().__init__()
        from reportlab.lib.utils import ImageReader
        reader = ImageReader(io.BytesIO(img_bytes))
        w, h = reader.getSize()
        scale = max_height / h if h > 0 else 1
        self.img_width = w * scale
        self.img_height = max_height
        self.img_bytes = img_bytes

    def wrap(self, availWidth, availHeight):
        return self.img_width, self.img_height

    def draw(self):
        self.canv.drawImage(
            io.BytesIO(self.img_bytes),
            0, 0,
            width=self.img_width,
            height=self.img_height,
            mask="auto",
        )


def _build_pdf_text_with_latex(text: str, style, story: list):
    """Add text with rendered LaTeX images to PDF story."""
    parts = split_text_and_latex(text)
    has_latex = any(p["type"] == "latex" for p in parts)

    if not has_latex:
        story.append(Paragraph(text, style))
        return

    # For mixed content: add text paragraphs and latex as images
    for part in parts:
        if part["type"] == "text":
            t = part["content"].strip()
            if t:
                story.append(Paragraph(t, style))
        else:
            try:
                fontsize = 16 if part.get("display") else 14
                img_bytes = latex_to_image(part["content"], fontsize=fontsize)
                img = Image(io.BytesIO(img_bytes))
                # Scale to reasonable size
                max_w = 160 * mm if part.get("display") else 80 * mm
                if img.drawWidth > max_w:
                    scale = max_w / img.drawWidth
                    img.drawWidth *= scale
                    img.drawHeight *= scale
                story.append(img)
            except Exception:
                # Fallback: raw LaTeX text
                story.append(Paragraph(f"${part['content']}$", style))


def export_pdf(variant_data: dict) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("VTitle", parent=styles["Title"], fontSize=16)
    q_style = ParagraphStyle("QStyle", parent=styles["Normal"], fontSize=11, spaceAfter=4)
    opt_style = ParagraphStyle("OptStyle", parent=styles["Normal"], fontSize=10, leftIndent=20)

    story = []
    v = variant_data["variant"]
    story.append(Paragraph(v.title, title_style))
    story.append(Paragraph(f"Fənn: {v.subject} | Sual sayı: {v.total_questions}", styles["Normal"]))
    story.append(Spacer(1, 12))

    for item in variant_data["questions"]:
        q = item["question"]
        story.append(Paragraph(f"<b>{item['order']}.</b>", q_style))
        _build_pdf_text_with_latex(q.question_text, q_style, story)
        if q.options:
            for key, val in q.options.items():
                _build_pdf_text_with_latex(f"{key}) {val}", opt_style, story)
        story.append(Spacer(1, 8))

    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>Cavablar</b>", styles["Heading2"]))
    for item in variant_data["questions"]:
        q = item["question"]
        _build_pdf_text_with_latex(f"{item['order']}. {q.correct_answer}", styles["Normal"], story)

    doc.build(story)
    return buffer.getvalue()


def _add_word_text_with_latex(doc: Document, text: str, bold_prefix: str = ""):
    """Add a paragraph with rendered LaTeX images to Word document."""
    parts = split_text_and_latex(text)
    has_latex = any(p["type"] == "latex" for p in parts)

    if not has_latex:
        p = doc.add_paragraph()
        if bold_prefix:
            run = p.add_run(bold_prefix)
            run.bold = True
        p.add_run(text)
        return

    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True

    for part in parts:
        if part["type"] == "text":
            p.add_run(part["content"])
        else:
            try:
                fontsize = 16 if part.get("display") else 14
                img_bytes = latex_to_image(part["content"], fontsize=fontsize)
                run = p.add_run()
                run.add_picture(io.BytesIO(img_bytes), height=Mm(5))
            except Exception:
                p.add_run(f"${part['content']}$")


def export_word(variant_data: dict) -> bytes:
    doc = Document()
    v = variant_data["variant"]

    doc.add_heading(v.title, level=1)
    doc.add_paragraph(f"Fənn: {v.subject} | Sual sayı: {v.total_questions}")

    for item in variant_data["questions"]:
        q = item["question"]
        _add_word_text_with_latex(doc, q.question_text, bold_prefix=f"{item['order']}. ")
        if q.options:
            for key, val in q.options.items():
                _add_word_text_with_latex(doc, val, bold_prefix=f"   {key}) ")

    doc.add_heading("Cavablar", level=2)
    for item in variant_data["questions"]:
        q = item["question"]
        _add_word_text_with_latex(doc, q.correct_answer, bold_prefix=f"{item['order']}. ")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
